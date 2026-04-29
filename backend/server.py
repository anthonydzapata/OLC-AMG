from fastapi import FastAPI, APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import json
import re
import io
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone
from emergentintegrations.llm.chat import LlmChat, UserMessage, ImageContent
from docx import Document
from docx.shared import Pt, RGBColor, Inches


ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")


# Define Models
class StatusCheck(BaseModel):
    model_config = ConfigDict(extra="ignore")  # Ignore MongoDB's _id field
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_name: str
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class StatusCheckCreate(BaseModel):
    client_name: str

# Add your routes to the router instead of directly to app
@api_router.get("/")
async def root():
    return {"message": "Hello World"}

@api_router.post("/status", response_model=StatusCheck)
async def create_status_check(input: StatusCheckCreate):
    status_dict = input.model_dump()
    status_obj = StatusCheck(**status_dict)
    
    # Convert to dict and serialize datetime to ISO string for MongoDB
    doc = status_obj.model_dump()
    doc['timestamp'] = doc['timestamp'].isoformat()
    
    _ = await db.status_checks.insert_one(doc)
    return status_obj

@api_router.get("/status", response_model=List[StatusCheck])
async def get_status_checks():
    # Exclude MongoDB's _id field from the query results
    status_checks = await db.status_checks.find({}, {"_id": 0}).to_list(1000)
    
    # Convert ISO string timestamps back to datetime objects
    for check in status_checks:
        if isinstance(check['timestamp'], str):
            check['timestamp'] = datetime.fromisoformat(check['timestamp'])
    
    return status_checks


# ─── BRIEF ANALYSIS (Claude via Emergent LLM Key) ─────────────────────────────

class AnalyzeBriefRequest(BaseModel):
    prompt: str

class AnalyzeBriefResponse(BaseModel):
    text: str

@api_router.post("/analyze-brief", response_model=AnalyzeBriefResponse)
async def analyze_brief(payload: AnalyzeBriefRequest):
    api_key = os.environ.get('EMERGENT_LLM_KEY')
    if not api_key:
        raise HTTPException(status_code=500, detail="LLM key not configured")
    if not payload.prompt or not payload.prompt.strip():
        raise HTTPException(status_code=400, detail="Prompt is required")

    try:
        chat = LlmChat(
            api_key=api_key,
            session_id=f"olc-brief-{uuid.uuid4()}",
            system_message="You are the internal creative intelligence for The Old Line Company (OLC)."
        ).with_model("anthropic", "claude-4-sonnet-20250514")

        response = await chat.send_message(UserMessage(text=payload.prompt))
        return AnalyzeBriefResponse(text=str(response))
    except Exception as e:
        logger.exception("Brief analysis failed")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")


# ─── INSPO ANALYZER (Gemini 2.5 Flash vision) ─────────────────────────────────

class AnalyzeImageRequest(BaseModel):
    image_base64: str  # raw base64 (no data: prefix) OR full data URL
    mime_type: Optional[str] = "image/jpeg"
    notes: Optional[str] = ""  # any creator notes to bias the analysis


INSPO_SYSTEM = (
    "You are a STRUCTURAL DESIGN ANALYST for The Old Line Company. Your output is "
    "consumed downstream by AI assistants that have aggressive content/IP guardrails. "
    "To prevent those guardrails from misreading your work as imitation of any specific "
    "creator, you operate under a hard rule: describe ONLY activity-level mechanics — "
    "geometry, ratios, color values, type properties, lighting setup, technique steps, "
    "equipment categories. NEVER name a designer, photographer, illustrator, director, "
    "studio, firm, brand, franchise, school, movement, decade, era, or country-of-origin. "
    "If a recognizable reference would be useful, translate it instead into its underlying "
    "structural grammar. Your job is to describe HOW, never WHO."
)

INSPO_PROMPT_TEMPLATE = """Analyze the attached reference image and return STRICT JSON only — no markdown, no commentary.

HARD RULES (violations are unacceptable):
- DO NOT name any designer, illustrator, photographer, director, studio, firm, brand, franchise, IP, school, movement, decade, era, or country-of-origin shorthand.
- DO NOT use phrases like "in the style of", "reminiscent of", "Bauhaus", "Art Deco", "1970s", "Japanese", "mid-century", "Saul Bass-style", etc.
- DESCRIBE every observation in mechanical / structural / activity terms only:
  * composition: angles, ratios, grid systems, focal points, axis, balance, negative-space distribution
  * color: HSL ranges, value relationships, saturation behavior, palette structure (count + relationships, NOT named movements)
  * form/type: weight, axis, contrast ratio, kerning behavior, geometric vs humanist construction
  * technique: discrete steps, tools/categories (not brand names), surface treatments
  * equipment: sensor format category, focal length category, lighting setup, exposure behavior — NEVER brand names
- Each "option" is an alternative STRUCTURAL DIRECTION (different mechanics) the creator could take for a NEW piece, not a reference to anyone's prior work.
- Be specific and measurable wherever possible. "30° rotation, 1:1.6 aspect, halftone at ~35lpi" beats "vintage poster vibe" every time.

For each dimension provide:
- "observed": a concise description in pure mechanics (1-3 sentences). State numbers, ratios, geometry.
- "options": THREE distinct alternative structural directions (under 25 words each).

Field set:

{
  "subject": {"observed": "...", "options": ["...", "...", "..."]},
  "method": {"observed": "technique mechanics: medium, surface, sequence of operations", "options": ["...", "...", "..."]},
  "composition": {"observed": "geometry: axis, ratios, focal points, negative space distribution, balance", "options": ["...", "...", "..."]},
  "structural_principles": {"observed": "underlying compositional/aesthetic principles in pure mechanics — what RULES are at work (e.g. 'symmetric one-point perspective with rule-of-thirds focal point; high-contrast monochrome with single accent hue; geometric sans-serif at consistent weight')", "philosophy": "1-2 sentence read on the mechanical philosophy: what is the work TRYING to accomplish structurally?", "options": ["...", "...", "..."]},
  "method_to_achieve": {"observed": "step-by-step mechanics: how to physically/digitally produce this look — tools by category, sequence, settings", "options": ["...", "...", "..."]},
  "technical_specs": {"observed": "equipment by CATEGORY only (e.g. 'medium-format digital sensor', 'tungsten-balanced 3200K lighting', 'wide-angle 24-35mm equivalent', 'screen-printing with 2-color separation') — never brand names or model numbers", "options": ["...", "...", "..."]}
}

Return that EXACT JSON shape and nothing else.

Creator notes (bias the analysis if present): {notes}
"""


def _extract_json(text: str) -> Dict[str, Any]:
    """Extract the first JSON object from a string, tolerating code fences."""
    # Strip code fences if present
    fenced = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.DOTALL)
    if fenced:
        text = fenced.group(1)
    # Find first {...} block
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if not match:
        raise ValueError("No JSON object found in model response")
    return json.loads(match.group(0))


@api_router.post("/analyze-image")
async def analyze_image(payload: AnalyzeImageRequest):
    api_key = os.environ.get('EMERGENT_LLM_KEY')
    if not api_key:
        raise HTTPException(status_code=500, detail="LLM key not configured")
    if not payload.image_base64:
        raise HTTPException(status_code=400, detail="image_base64 is required")

    # Strip data URL prefix if present
    b64 = payload.image_base64
    if b64.startswith('data:'):
        comma = b64.find(',')
        if comma == -1:
            raise HTTPException(status_code=400, detail="Malformed data URL")
        b64 = b64[comma + 1:]

    try:
        chat = LlmChat(
            api_key=api_key,
            session_id=f"olc-inspo-{uuid.uuid4()}",
            system_message=INSPO_SYSTEM,
        ).with_model("gemini", "gemini-2.5-flash")

        prompt = INSPO_PROMPT_TEMPLATE.replace("{notes}", payload.notes or "(none)")

        image = ImageContent(image_base64=b64)
        response = await chat.send_message(UserMessage(
            text=prompt,
            file_contents=[image]
        ))

        raw = str(response)
        try:
            data = _extract_json(raw)
        except Exception:
            # Return raw so frontend can show it for manual triage instead of failing silently
            return {"raw": raw, "parsed": None}

        return {"raw": raw, "parsed": data}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Image analysis failed")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")


# ─── STRUCTURAL PROMPT GENERATION ─────────────────────────────────────────────

class GeneratePromptRequest(BaseModel):
    title: Optional[str] = ""
    subject: Optional[str] = ""
    method: Optional[str] = ""
    composition: Optional[str] = ""
    structural_principles: Optional[str] = ""
    method_to_achieve: Optional[str] = ""
    technical_specs: Optional[str] = ""
    notes: Optional[str] = ""

@api_router.post("/generate-prompt")
async def generate_prompt(payload: GeneratePromptRequest):
    """Assemble a structured executor prompt from the chosen directions.
    Pure template — no LLM call, no credit cost. Output is engineered to pass
    downstream guardrails: zero attribution, pure mechanics, explicit creative
    freedom WITHIN a narrow structural constraint set.
    """
    def _s(v):
        return (v or "").strip() if isinstance(v, str) else ""

    title = _s(payload.title) or "Untitled directive"
    lines = [
        f"# OPERATING DIRECTIVE — {title}",
        "",
        "You are an executor producing a NEW, ORIGINAL piece. The structure below is a "
        "STRUCTURAL CONSTRAINT SET — composition geometry, color systems, type/form properties, "
        "technique mechanics. Within this structure an infinite number of valid choices exist. "
        "MAKE YOUR OWN CHOICES.",
        "",
        "READ THIS CAREFULLY:",
        "- The constraints describe HOW (mechanics), not WHO (no creator, firm, era is referenced).",
        "- Do not interpret any line below as 'in the style of' anyone or anything.",
        "- Your output must be entirely original. The constraints define the structural grammar; "
        "you create freely within it.",
        "- If a constraint feels under-defined, fill it with your own original choice. That is the point.",
        "",
        "## SUBJECT",
        _s(payload.subject) or "(creator's choice within OLC voice)",
        "",
        "## METHOD (technique mechanics)",
        _s(payload.method) or "(creator's choice)",
        "",
        "## COMPOSITION (geometry / axis / ratios / focal behavior)",
        _s(payload.composition) or "(creator's choice)",
        "",
        "## STRUCTURAL & COMPOSITIONAL PRINCIPLES",
        _s(payload.structural_principles) or "(creator's choice)",
        "",
        "## METHOD TO ACHIEVE THE LOOK (sequence of operations)",
        _s(payload.method_to_achieve) or "(creator's choice)",
        "",
        "## TECHNICAL SPECS / SETTINGS (categories, never brand names)",
        _s(payload.technical_specs) or "(creator's choice)",
        "",
        "## CREATOR NOTES",
        _s(payload.notes) or "(none)",
        "",
        "## DELIVERY RULES",
        "- Honor the structural constraints as a grammar, not a costume.",
        "- Make and disclose your specific choices for any element that was creator's-choice.",
        "- Do not name or invoke any prior creator, brand, or work in your output or rationale.",
        "- Surface the choices you made and the mechanical reasoning behind them.",
    ]
    return {"prompt": "\n".join(lines)}


# ─── EXECUTE: GENERATE VISUAL DELIVERABLE (Gemini Nano Banana) ────────────────

class ExecuteVisualRequest(BaseModel):
    prompt: str
    refine: Optional[str] = ""  # optional iteration note
    reference_image_b64: Optional[str] = None  # optional reference image (currently unused for first iteration; kept for forward compat)

@api_router.post("/execute-visual")
async def execute_visual(payload: ExecuteVisualRequest):
    api_key = os.environ.get('EMERGENT_LLM_KEY')
    if not api_key:
        raise HTTPException(status_code=500, detail="LLM key not configured")
    if not (payload.prompt or "").strip():
        raise HTTPException(status_code=400, detail="prompt is required")

    full_prompt = payload.prompt.strip()
    if payload.refine and payload.refine.strip():
        full_prompt += "\n\n## ITERATION NOTE\n" + payload.refine.strip()

    try:
        chat = LlmChat(
            api_key=api_key,
            session_id=f"olc-visual-{uuid.uuid4()}",
            system_message="You are a visual executor producing original imagery from structural directives.",
        ).with_model("gemini", "gemini-3.1-flash-image-preview").with_params(modalities=["image", "text"])

        msg = UserMessage(text=full_prompt)
        text, images = await chat.send_message_multimodal_response(msg)

        if not images:
            return {"text": str(text or ""), "images": []}

        out_images = []
        for img in images:
            out_images.append({
                "mime_type": img.get("mime_type", "image/png"),
                "data": img.get("data", "")
            })
        return {"text": str(text or ""), "images": out_images}
    except Exception as e:
        logger.exception("Execute visual failed")
        raise HTTPException(status_code=500, detail=f"Visual generation failed: {str(e)}")


# ─── EXECUTE: GENERATE TEXT TREATMENT (Claude Sonnet) ─────────────────────────

class ExecuteTreatmentRequest(BaseModel):
    prompt: str
    refine: Optional[str] = ""

TREATMENT_SYSTEM = (
    "You are a senior creative director for The Old Line Company. You receive a "
    "structural directive (mechanics only) and you produce a polished, original creative "
    "treatment / script / copy / shot list as the deliverable. You operate under a strict rule: "
    "describe everything in original, mechanical, structural language. Do not name any "
    "designer, photographer, director, illustrator, studio, firm, brand, franchise, "
    "school, movement, decade, era, or country-of-origin. The directive describes HOW, never WHO. "
    "Your treatment must be entirely original."
)

@api_router.post("/execute-treatment")
async def execute_treatment(payload: ExecuteTreatmentRequest):
    api_key = os.environ.get('EMERGENT_LLM_KEY')
    if not api_key:
        raise HTTPException(status_code=500, detail="LLM key not configured")
    if not (payload.prompt or "").strip():
        raise HTTPException(status_code=400, detail="prompt is required")

    full_prompt = payload.prompt.strip()
    if payload.refine and payload.refine.strip():
        full_prompt += "\n\n## ITERATION NOTE\n" + payload.refine.strip()
    full_prompt += (
        "\n\n## YOUR TASK\n"
        "Produce the deliverable that fulfills the directive above. Format appropriately for the "
        "implied medium (poster spec, video shot list, copy block, script, etc.). "
        "Use clear section headers. Disclose your specific choices for under-defined elements "
        "and the mechanical reasoning behind them. Do not name any creator, firm, era, or prior work."
    )

    try:
        chat = LlmChat(
            api_key=api_key,
            session_id=f"olc-treatment-{uuid.uuid4()}",
            system_message=TREATMENT_SYSTEM,
        ).with_model("anthropic", "claude-4-sonnet-20250514")

        response = await chat.send_message(UserMessage(text=full_prompt))
        return {"text": str(response)}
    except Exception as e:
        logger.exception("Execute treatment failed")
        raise HTTPException(status_code=500, detail=f"Treatment generation failed: {str(e)}")


# ─── DOCX EXPORT ──────────────────────────────────────────────────────────────

class ExportDocxRequest(BaseModel):
    title: Optional[str] = "OLC Inspo Analysis"
    image_base64: Optional[str] = None  # optional — embed reference image
    mime_type: Optional[str] = "image/jpeg"
    breakdown: Dict[str, Any] = Field(default_factory=dict)  # observed values
    choices: Dict[str, Any] = Field(default_factory=dict)    # chosen values
    notes: Optional[str] = ""
    generated_prompt: Optional[str] = ""


def _add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    return h


@api_router.post("/export-docx")
async def export_docx(payload: ExportDocxRequest):
    doc = Document()

    # Title
    title = doc.add_heading(payload.title or "OLC Inspo Analysis", level=0)
    sub = doc.add_paragraph()
    sub_run = sub.add_run(f"The Old Line Company  •  {datetime.now().strftime('%B %d, %Y')}")
    sub_run.italic = True
    sub_run.font.size = Pt(10)

    # Embed reference image if provided
    if payload.image_base64:
        try:
            import base64 as b64lib
            b64 = payload.image_base64
            if b64.startswith('data:'):
                comma = b64.find(',')
                if comma != -1:
                    b64 = b64[comma + 1:]
            img_bytes = b64lib.b64decode(b64)
            img_stream = io.BytesIO(img_bytes)
            doc.add_heading("Reference", level=1)
            doc.add_picture(img_stream, width=Inches(4.0))
        except Exception:
            logger.warning("Could not embed reference image into docx", exc_info=True)

    # Breakdown sections
    sections = [
        ("subject", "Subject"),
        ("method", "Method (Technique Mechanics)"),
        ("composition", "Composition (Geometry & Axis)"),
        ("structural_principles", "Structural & Compositional Principles"),
        ("method_to_achieve", "Method To Achieve The Look"),
        ("technical_specs", "Technical Specs / Settings"),
    ]

    doc.add_heading("Breakdown & Direction", level=1)
    for key, label in sections:
        doc.add_heading(label, level=2)
        observed = (payload.breakdown.get(key) or "").strip() or "—"
        chosen = (payload.choices.get(key) or "").strip() or "—"
        p = doc.add_paragraph()
        r = p.add_run("Observed: ")
        r.bold = True
        p.add_run(observed)
        p2 = doc.add_paragraph()
        r2 = p2.add_run("Direction chosen: ")
        r2.bold = True
        p2.add_run(chosen)

    # Notes
    if payload.notes and payload.notes.strip():
        doc.add_heading("Work Product Notes", level=1)
        doc.add_paragraph(payload.notes.strip())

    # Generated prompt
    if payload.generated_prompt and payload.generated_prompt.strip():
        doc.add_heading("Brief Prompt (for executor)", level=1)
        for line in payload.generated_prompt.split('\n'):
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.strip().startswith('- '):
                doc.add_paragraph(line.strip()[2:], style='List Bullet')
            elif line.strip() == '':
                doc.add_paragraph('')
            else:
                doc.add_paragraph(line)

    # Stream back
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = re.sub(r"[^A-Za-z0-9_-]+", "_", (payload.title or "OLC_Inspo_Analysis"))[:60]
    filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
